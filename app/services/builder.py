from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def build_translated_docx(
    source_path: Path,
    output_path: Path,
    translations: dict[str, str],
) -> None:
    document = Document(source_path)

    for index, paragraph in enumerate(document.paragraphs, start=1):
        _replace_paragraph_text(paragraph, f"p{index}", translations)

    table_index = 0
    for table in document.tables:
        table_index += 1
        _replace_table_text(table, table_index, translations)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _replace_table_text(
    table: Table,
    table_index: int,
    translations: dict[str, str],
) -> None:
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            cell_prefix = f"t{table_index}r{row_index}c{cell_index}"

            for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                block_id = f"{cell_prefix}p{paragraph_index}"
                _replace_paragraph_text(paragraph, block_id, translations)


def _replace_paragraph_text(
    paragraph: Paragraph,
    block_id: str,
    translations: dict[str, str],
) -> None:
    translated_text = translations.get(block_id)
    if translated_text is None:
        return

    if not paragraph.runs:
        paragraph.add_run(translated_text)
        return

    paragraph.runs[0].text = translated_text
    for run in paragraph.runs[1:]:
        run.text = ""
