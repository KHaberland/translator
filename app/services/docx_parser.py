import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.models.schemas import DocumentBlock


TECHNICAL_TEXT_RE = re.compile(
    r"^(?:[A-Z]{2,}\s?\d+(?:[-:.]\d+)*|v\d+(?:\.\d+){1,3}|[A-Z0-9_.-]{2,})$"
)


def extract_docx_blocks(file_path: Path) -> list[DocumentBlock]:
    document = Document(file_path)
    blocks: list[DocumentBlock] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        _append_paragraph_block(blocks, paragraph, f"p{index}", f"paragraph:{index}")

    table_index = 0
    for table in document.tables:
        table_index += 1
        _append_table_blocks(blocks, table, table_index)

    return blocks


def _append_table_blocks(blocks: list[DocumentBlock], table: Table, table_index: int) -> None:
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            cell_prefix = f"t{table_index}r{row_index}c{cell_index}"
            location_prefix = f"table:{table_index}:row:{row_index}:cell:{cell_index}"

            for paragraph_index, paragraph in enumerate(_iter_cell_paragraphs(cell), start=1):
                block_id = f"{cell_prefix}p{paragraph_index}"
                location = f"{location_prefix}:paragraph:{paragraph_index}"
                _append_paragraph_block(blocks, paragraph, block_id, location)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    yield from cell.paragraphs


def _append_paragraph_block(
    blocks: list[DocumentBlock],
    paragraph: Paragraph,
    block_id: str,
    location: str,
) -> None:
    text = paragraph.text.strip()
    if not text:
        return

    is_translatable = not _looks_non_translatable(text, paragraph)
    blocks.append(
        DocumentBlock(
            block_id=block_id,
            text=text,
            location=location,
            translatable=is_translatable,
            reason=None if is_translatable else "technical_or_code_like_text",
        )
    )


def _looks_non_translatable(text: str, paragraph: Paragraph) -> bool:
    normalized = " ".join(text.split())
    if _has_code_like_style(paragraph):
        return True

    if TECHNICAL_TEXT_RE.fullmatch(normalized):
        return True

    return _looks_like_code(normalized)


def _has_code_like_style(paragraph: Paragraph) -> bool:
    style_name = getattr(paragraph.style, "name", "") or ""
    normalized_style = style_name.lower()
    return "code" in normalized_style or "formula" in normalized_style


def _looks_like_code(text: str) -> bool:
    code_markers = ("{", "}", "</", "/>", "==", "!=", "=>", "::", "```")
    if any(marker in text for marker in code_markers):
        return True

    return bool(re.fullmatch(r"[\w.-]+\([^)]*\)", text))
