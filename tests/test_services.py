import asyncio
import zipfile
from types import SimpleNamespace

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi import HTTPException
from fastapi.testclient import TestClient
from redis.exceptions import RedisError

from app.api.download import get_translation_download
from app.api.estimate import estimate_docx
from app.api.status import get_translation_status
from app.api.translate import translate_docx
from app.core.cache import RedisTranslationCache, build_cache_key
from app.core.config import Settings
from app.core.job_store import InMemoryJobStore
from app.core.progress_events import build_progress_event
from app.main import app
from app.models.jobs import JobStatus, TranslationJob
from app.models.schemas import DocumentBlock, LanguageCode
from app.services.builder import build_translated_docx
from app.services.cost_estimator import estimate_translation_cost
from app.services.docx_parser import extract_docx_blocks
from app.services.price_estimator import (
    budget_status,
    estimate_output_tokens,
    estimate_translation_cost_usd,
)
from app.services.segmenter import build_translation_batches
from app.services.translator import DocumentProcessingError, translate_docx_file
from app.services.translation_cache import TranslationCache, normalize_translation_cache_text
from app.services.translation_memory import SQLiteTranslationMemory
from workers.translation_worker import process_translation_job


def test_extract_docx_blocks_from_paragraphs_and_tables(tmp_path):
    source_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Document title")
    document.add_paragraph("")
    document.add_paragraph("API123")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text"
    document.save(source_path)

    blocks = extract_docx_blocks(source_path)

    assert [block.block_id for block in blocks] == ["p1", "p3", "t1r1c1p1"]
    assert blocks[0].text == "Document title"
    assert blocks[1].translatable is False
    assert blocks[2].location == "table:1:row:1:cell:1:paragraph:1"


def test_build_translated_docx_replaces_paragraphs_and_table_cells(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("Hello")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "World"
    document.save(source_path)

    build_translated_docx(
        source_path,
        output_path,
        {
            "p1": "Привет",
            "t1r1c1p1": "Мир",
        },
    )

    result = Document(output_path)
    assert result.paragraphs[0].text == "Привет"
    assert result.tables[0].cell(0, 0).text == "Мир"


def test_build_translated_docx_preserves_mixed_run_formatting(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    bold_run = paragraph.add_run("Hello ")
    bold_run.bold = True
    italic_run = paragraph.add_run("world")
    italic_run.italic = True
    document.save(source_path)

    build_translated_docx(source_path, output_path, {"p1": "Привет мир"})

    result_paragraph = Document(output_path).paragraphs[0]
    assert result_paragraph.text == "Привет мир"
    assert len(result_paragraph.runs) == 2
    assert result_paragraph.runs[0].bold is True
    assert result_paragraph.runs[1].italic is True


def test_build_translated_docx_preserves_list_style(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("First item", style="List Bullet")
    document.save(source_path)

    build_translated_docx(source_path, output_path, {"p1": "Первый пункт"})

    result_paragraph = Document(output_path).paragraphs[0]
    assert result_paragraph.text == "Первый пункт"
    assert result_paragraph.style.name == "List Bullet"


def test_build_translated_docx_preserves_table_structure(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "One"
    table.cell(0, 1).text = "Two"
    table.cell(1, 0).text = "Three"
    table.cell(1, 1).text = "Four"
    document.save(source_path)

    build_translated_docx(
        source_path,
        output_path,
        {
            "t1r1c1p1": "Один",
            "t1r2c2p1": "Четыре",
        },
    )

    result_table = Document(output_path).tables[0]
    assert len(result_table.rows) == 2
    assert len(result_table.columns) == 2
    assert result_table.cell(0, 0).text == "Один"
    assert result_table.cell(0, 1).text == "Two"
    assert result_table.cell(1, 1).text == "Четыре"


def test_build_translated_docx_preserves_hyperlink_container(tmp_path):
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph("Read ")
    _add_hyperlink(paragraph, "docs", "https://example.com")
    document.save(source_path)

    build_translated_docx(source_path, output_path, {"p1": "Читайте документацию"})

    result_paragraph = Document(output_path).paragraphs[0]
    assert result_paragraph.text == "Читайте документацию"

    with zipfile.ZipFile(output_path) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
        relationships_xml = docx_zip.read("word/_rels/document.xml.rels").decode("utf-8")

    assert "<w:hyperlink" in document_xml
    assert "https://example.com" in relationships_xml


def test_translation_cache_normalizes_repeated_text():
    cache = TranslationCache(LanguageCode.EN, LanguageCode.RU)
    first_block = _block("b1", " Repeated   text ")
    second_block = _block("b2", "Repeated text")

    assert normalize_translation_cache_text(" Repeated   text ") == "Repeated text"
    assert cache.original_block_id_for(first_block) is None
    assert cache.original_block_id_for(second_block) == "b1"

    cache.remember_translation(first_block.text, "Повторяющийся текст")
    assert cache.translation_for(second_block.text) == "Повторяющийся текст"


def test_redis_translation_cache_uses_ttl_and_normalized_key():
    redis_client = _FakeRedisClient()
    settings = Settings(translation_cache_ttl_seconds=60)
    cache = RedisTranslationCache(settings, redis_client=redis_client)
    key = build_cache_key(" Repeated   text ", LanguageCode.EN, LanguageCode.RU)

    assert key == build_cache_key("Repeated text", LanguageCode.EN, LanguageCode.RU)

    cache.set_translation(key, "Повторяющийся текст")

    assert redis_client.set_calls == [(key, 60, "Повторяющийся текст")]
    assert cache.get_translation(key) == "Повторяющийся текст"


def test_redis_translation_cache_falls_back_when_redis_is_unavailable():
    cache = RedisTranslationCache(
        Settings(translation_cache_ttl_seconds=60),
        redis_client=_FailingRedisClient(),
    )

    assert cache.get_translation("translation_cache:missing") is None
    cache.set_translation("translation_cache:missing", "ignored")


def test_segmenter_respects_batch_limits_and_skips_non_translatable_blocks():
    blocks = [
        _block("b1", "12345"),
        _block("b2", "1234"),
        _block("b3", "skip", translatable=False),
        _block("b4", "12"),
    ]

    batches = build_translation_batches(blocks, max_batch_chars=6, max_batch_blocks=2)

    assert [[block.block_id for block in batch] for batch in batches] == [
        ["b1"],
        ["b2", "b4"],
    ]


def test_segmenter_rejects_single_block_over_char_limit():
    with pytest.raises(ValueError, match="exceeds max_batch_chars"):
        build_translation_batches(
            [_block("b1", "too long")],
            max_batch_chars=3,
            max_batch_blocks=10,
        )


def test_cost_estimator_counts_only_translatable_blocks():
    estimate = estimate_translation_cost(
        [
            _block("b1", "12345678"),
            _block("b2", "ignored", translatable=False),
        ]
    )

    assert estimate.translatable_characters == 8
    assert estimate.estimated_tokens == 2


def test_price_estimator_calculates_output_tokens_cost_and_budget():
    settings = Settings(
        openai_input_price_per_1m_tokens=0.15,
        openai_output_price_per_1m_tokens=0.60,
    )
    output_tokens = estimate_output_tokens(10, 1.2)

    assert output_tokens == 12
    assert estimate_translation_cost_usd(10, output_tokens, settings) == 0.000009
    assert budget_status(0.01, 1) == "ok"
    assert budget_status(1.01, 1) == "exceeded"


def test_estimate_endpoint_accepts_docx_and_does_not_enqueue_job(tmp_path, monkeypatch):
    docx_path = tmp_path / "estimate.docx"
    document = Document()
    document.add_paragraph("Hello world")
    document.add_paragraph("API123")
    document.save(docx_path)
    file = _UploadFileStub(
        filename="estimate.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=docx_path.read_bytes(),
    )

    monkeypatch.setattr(
        "app.api.estimate.get_settings",
        lambda: Settings(
            tmp_dir=tmp_path / "tmp",
            translation_budget_usd=10,
            estimated_output_token_multiplier=1.2,
            translation_memory_db_path=tmp_path / "memory.sqlite3",
        ),
    )
    monkeypatch.setattr(
        "app.api.estimate.extract_docx_blocks",
        lambda path: [
            _block("b1", "Hello world"),
            _block("b2", "API123", translatable=False),
        ],
    )
    monkeypatch.setattr(
        "app.api.estimate.get_job_store",
        lambda: pytest.fail("estimate should not create jobs"),
        raising=False,
    )
    monkeypatch.setattr(
        "app.api.estimate._enqueue_translation_job",
        lambda job_id: pytest.fail("estimate should not enqueue Celery tasks"),
        raising=False,
    )

    response = asyncio.run(
        estimate_docx(
            file=file,
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
        )
    )

    assert response.file_name == "estimate.docx"
    assert response.translatable_blocks == 1
    assert response.skipped_blocks == 1
    assert response.estimated_characters == 11
    assert response.estimated_input_tokens == 3
    assert response.estimated_output_tokens == 4
    assert response.estimated_total_tokens == 7
    assert response.estimated_cost_usd > 0
    assert response.budget_status == "ok"


def test_estimate_endpoint_returns_exceeded_budget(tmp_path, monkeypatch):
    docx_path = _create_docx(tmp_path, "expensive.docx", "Hello")
    file = _UploadFileStub(
        filename="expensive.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=docx_path.read_bytes(),
    )

    monkeypatch.setattr(
        "app.api.estimate.get_settings",
        lambda: Settings(
            tmp_dir=tmp_path / "tmp",
            translation_budget_usd=0,
            estimated_output_token_multiplier=1.2,
            translation_memory_db_path=tmp_path / "memory.sqlite3",
        ),
    )

    response = asyncio.run(
        estimate_docx(
            file=file,
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
        )
    )

    assert response.estimated_characters == 5
    assert response.estimated_total_tokens > 0
    assert response.budget_status == "exceeded"


def test_translate_endpoint_rejects_equal_languages_before_reading_file():
    file = SimpleNamespace(filename="document.docx")

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(
            translate_docx(
                file=file,
                source_lang=LanguageCode.EN,
                target_lang=LanguageCode.EN,
            )
        )

    assert exc_info.value.status_code == 400
    assert exc_info.value.detail == "source_lang and target_lang must be different"


def test_translate_endpoint_rejects_non_docx_file():
    file = SimpleNamespace(filename="document.txt", content_type="text/plain")

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(
            translate_docx(
                file=file,
                source_lang=LanguageCode.EN,
                target_lang=LanguageCode.RU,
            )
        )

    assert exc_info.value.status_code == 400
    assert exc_info.value.detail == "only DOCX files are supported"


def test_translate_endpoint_rejects_invalid_docx_content_type():
    file = SimpleNamespace(filename="document.docx", content_type="text/plain")

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(
            translate_docx(
                file=file,
                source_lang=LanguageCode.EN,
                target_lang=LanguageCode.RU,
            )
        )

    assert exc_info.value.status_code == 400
    assert exc_info.value.detail == "invalid DOCX content type"


def test_translate_docx_file_uses_mock_ai_without_api_key(tmp_path):
    source_path = tmp_path / "source.docx"
    output_dir = tmp_path / "outputs"
    document = Document()
    document.add_paragraph("Hello")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "World"
    document.save(source_path)

    result = asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                mock_ai_enabled=True,
                output_dir=output_dir,
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                translation_memory_db_path=tmp_path / "memory.sqlite3",
            ),
        )
    )

    translated_document = Document(result.file_path)
    assert result.status == "completed"
    assert result.file_name == "source_translated_to_ru.docx"
    assert translated_document.paragraphs[0].text == "Hello [ru]"
    assert translated_document.tables[0].cell(0, 0).text == "World [ru]"


def test_translate_docx_file_uses_redis_cache_hit_without_ai_call(tmp_path, monkeypatch):
    source_path = _create_docx(tmp_path, "source.docx", "Hello")
    output_dir = tmp_path / "outputs"
    cache = _PipelineCache({"Hello": "Привет"})

    monkeypatch.setattr("app.services.translator.get_translation_cache", lambda settings: cache)
    monkeypatch.setattr(
        "app.services.translator.get_translation_client",
        lambda settings: pytest.fail("AI client should not be created on cache hit"),
    )

    result = asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                output_dir=output_dir,
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                translation_memory_db_path=tmp_path / "memory.sqlite3",
            ),
        )
    )

    translated_document = Document(result.file_path)
    assert result.estimated_characters == 0
    assert translated_document.paragraphs[0].text == "Привет"
    assert cache.stored == []


def test_translate_docx_file_saves_redis_cache_after_ai_response(tmp_path, monkeypatch):
    source_path = _create_docx(tmp_path, "source.docx", "Hello")
    cache = _PipelineCache({})
    client = _FakeTranslationClient({"p1": "Привет"})

    monkeypatch.setattr("app.services.translator.get_translation_cache", lambda settings: cache)
    monkeypatch.setattr("app.services.translator.get_translation_client", lambda settings: client)

    result = asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                output_dir=tmp_path / "outputs",
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                translation_memory_db_path=tmp_path / "memory.sqlite3",
            ),
        )
    )

    translated_document = Document(result.file_path)
    assert client.requested_block_ids == [["p1"]]
    assert translated_document.paragraphs[0].text == "Привет"
    assert cache.stored == [
        (
            build_cache_key("Hello", LanguageCode.EN, LanguageCode.RU),
            "Привет",
        )
    ]


def test_translate_docx_file_uses_translation_memory_hit_without_ai_call(tmp_path, monkeypatch):
    source_path = _create_docx(tmp_path, "source.docx", " Hello   world ")
    cache = _PipelineCache({})
    memory = SQLiteTranslationMemory(tmp_path / "memory.sqlite3")
    memory.save_translation(
        "Hello world",
        "Привет, мир",
        LanguageCode.EN,
        LanguageCode.RU,
    )

    monkeypatch.setattr("app.services.translator.get_translation_cache", lambda settings: cache)
    monkeypatch.setattr(
        "app.services.translator.get_translation_client",
        lambda settings: pytest.fail("AI client should not be created on memory hit"),
    )

    result = asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                output_dir=tmp_path / "outputs",
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                translation_memory_db_path=tmp_path / "memory.sqlite3",
            ),
        )
    )

    translated_document = Document(result.file_path)
    assert result.estimated_characters == 0
    assert translated_document.paragraphs[0].text == "Привет, мир"
    assert memory.frequency_for("Hello world", LanguageCode.EN, LanguageCode.RU) == 2
    assert cache.stored == [
        (
            build_cache_key(" Hello   world ", LanguageCode.EN, LanguageCode.RU),
            "Привет, мир",
        )
    ]


def test_translate_docx_file_saves_translation_memory_after_ai_response(tmp_path, monkeypatch):
    source_path = _create_docx(tmp_path, "source.docx", "Hello")
    cache = _PipelineCache({})
    client = _FakeTranslationClient({"p1": "Привет"})
    memory_path = tmp_path / "memory.sqlite3"

    monkeypatch.setattr("app.services.translator.get_translation_cache", lambda settings: cache)
    monkeypatch.setattr("app.services.translator.get_translation_client", lambda settings: client)

    asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                output_dir=tmp_path / "outputs",
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                translation_memory_db_path=memory_path,
            ),
        )
    )

    memory = SQLiteTranslationMemory(memory_path)
    assert memory.lookup_exact(" Hello ", LanguageCode.EN, LanguageCode.RU) == "Привет"
    assert memory.frequency_for("Hello", LanguageCode.EN, LanguageCode.RU) == 1


def test_translate_docx_file_sends_only_relevant_glossary_terms_per_batch(
    tmp_path,
    monkeypatch,
):
    source_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("WPS manual")
    document.add_paragraph("Plain paragraph")
    document.save(source_path)
    cache = _PipelineCache({})
    client = _FakeTranslationClient(
        {
            "p1": "Руководство по сварочному позиционеру",
            "p2": "Обычный абзац",
        }
    )
    memory_path = tmp_path / "memory.sqlite3"
    memory = SQLiteTranslationMemory(memory_path)
    memory.add_glossary_term(
        "WPS",
        "сварочный позиционер",
        LanguageCode.EN,
        LanguageCode.RU,
    )
    memory.add_glossary_term(
        "API",
        "программный интерфейс",
        LanguageCode.EN,
        LanguageCode.RU,
    )

    monkeypatch.setattr("app.services.translator.get_translation_cache", lambda settings: cache)
    monkeypatch.setattr("app.services.translator.get_translation_client", lambda settings: client)

    asyncio.run(
        translate_docx_file(
            source_path=source_path,
            original_filename="source.docx",
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
            settings=Settings(
                output_dir=tmp_path / "outputs",
                upload_dir=tmp_path / "uploads",
                tmp_dir=tmp_path / "tmp",
                max_batch_blocks=1,
                translation_memory_db_path=memory_path,
            ),
        )
    )

    assert client.requested_block_ids == [["p1"], ["p2"]]
    assert client.requested_glossary_terms == [
        [("WPS", "сварочный позиционер")],
        [],
    ]


def test_translate_endpoint_creates_queued_job_without_ai_call(tmp_path, monkeypatch):
    store = InMemoryJobStore()
    enqueued_job_ids: list[str] = []
    upload_dir = tmp_path / "uploads"
    docx_path = _create_docx(tmp_path, "queued.docx", "Hello")
    file = _UploadFileStub(
        filename="queued.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=docx_path.read_bytes(),
    )

    monkeypatch.setattr(
        "app.api.translate.get_settings",
        lambda: Settings(
            mock_ai_enabled=True,
            upload_dir=upload_dir,
            output_dir=tmp_path / "outputs",
            tmp_dir=tmp_path / "tmp",
            translation_memory_db_path=tmp_path / "memory.sqlite3",
        ),
    )
    monkeypatch.setattr("app.api.translate.get_job_store", lambda: store)
    monkeypatch.setattr(
        "app.api.translate._enqueue_translation_job",
        lambda job_id: enqueued_job_ids.append(job_id),
    )

    response = asyncio.run(
        translate_docx(
            file=file,
            source_lang=LanguageCode.EN,
            target_lang=LanguageCode.RU,
        )
    )

    stored_job = store.get_job(response.job_id)
    assert response.status == JobStatus.QUEUED
    assert stored_job is not None
    assert stored_job.status == JobStatus.QUEUED
    assert stored_job.progress == 0
    assert stored_job.original_filename == "queued.docx"
    assert stored_job.upload_path.startswith(upload_dir.as_posix())
    assert enqueued_job_ids == [response.job_id]


def test_status_endpoint_returns_job_state(monkeypatch):
    store = InMemoryJobStore()
    job = TranslationJob(
        job_id="job-1",
        source_lang=LanguageCode.EN,
        target_lang=LanguageCode.RU,
        original_filename="source.docx",
        upload_path="uploads/source.docx",
    )
    store.create_job(job)
    store.update_job(
        "job-1",
        status=JobStatus.TRANSLATING,
        progress=45,
    )
    monkeypatch.setattr("app.api.status.get_job_store", lambda: store)

    response = get_translation_status("job-1")

    assert response.job_id == "job-1"
    assert response.status == JobStatus.TRANSLATING
    assert response.progress == 45
    assert response.result_file is None
    assert response.error is None


def test_status_endpoint_returns_404_for_missing_job(monkeypatch):
    monkeypatch.setattr("app.api.status.get_job_store", lambda: InMemoryJobStore())

    with pytest.raises(HTTPException) as exc_info:
        get_translation_status("missing")

    assert exc_info.value.status_code == 404
    assert exc_info.value.detail == "translation job not found"


def test_download_endpoint_returns_completed_result(tmp_path, monkeypatch):
    store = InMemoryJobStore()
    result_path = _create_docx(tmp_path, "translated.docx", "Привет")
    store.create_job(
        _job(
            "job-download",
            status=JobStatus.COMPLETED,
            progress=100,
            result_file=result_path.as_posix(),
        )
    )
    monkeypatch.setattr("app.api.download.get_job_store", lambda: store)

    response = get_translation_download("job-download")

    assert response.path == result_path
    assert response.media_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.filename == "translated.docx"


def test_download_endpoint_returns_404_for_missing_job(monkeypatch):
    monkeypatch.setattr("app.api.download.get_job_store", lambda: InMemoryJobStore())

    with pytest.raises(HTTPException) as exc_info:
        get_translation_download("missing")

    assert exc_info.value.status_code == 404
    assert exc_info.value.detail == "translation job not found"


def test_download_endpoint_returns_409_when_result_is_not_ready(monkeypatch):
    store = InMemoryJobStore()
    store.create_job(_job("job-not-ready", status=JobStatus.TRANSLATING, progress=50))
    monkeypatch.setattr("app.api.download.get_job_store", lambda: store)

    with pytest.raises(HTTPException) as exc_info:
        get_translation_download("job-not-ready")

    assert exc_info.value.status_code == 409
    assert exc_info.value.detail == "translation result is not ready"


def test_download_endpoint_returns_404_when_result_file_is_missing(tmp_path, monkeypatch):
    store = InMemoryJobStore()
    store.create_job(
        _job(
            "job-missing-file",
            status=JobStatus.COMPLETED,
            progress=100,
            result_file=(tmp_path / "missing.docx").as_posix(),
        )
    )
    monkeypatch.setattr("app.api.download.get_job_store", lambda: store)

    with pytest.raises(HTTPException) as exc_info:
        get_translation_download("job-missing-file")

    assert exc_info.value.status_code == 404
    assert exc_info.value.detail == "translation result file not found"


def test_stream_endpoint_returns_sse_headers(monkeypatch):
    store = InMemoryJobStore()
    store.create_job(
        _job(
            "job-stream-headers",
            status=JobStatus.COMPLETED,
            progress=100,
        )
    )
    monkeypatch.setattr("app.api.stream.get_job_store", lambda: store)
    monkeypatch.setattr(
        "app.api.stream.get_progress_event_store",
        lambda: _StreamEventStore([]),
    )

    response = TestClient(app).get("/stream/job-stream-headers")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    assert response.headers["cache-control"] == "no-cache"


def test_stream_endpoint_sends_completed_final_event(monkeypatch):
    store = InMemoryJobStore()
    store.create_job(
        _job(
            "job-stream-completed",
            status=JobStatus.COMPLETED,
            progress=100,
        )
    )
    monkeypatch.setattr("app.api.stream.get_job_store", lambda: store)
    monkeypatch.setattr(
        "app.api.stream.get_progress_event_store",
        lambda: _StreamEventStore([]),
    )

    response = TestClient(app).get("/stream/job-stream-completed")

    assert response.status_code == 200
    assert "event: progress" in response.text
    assert '"stage":"completed"' in response.text
    assert '"progress":100' in response.text


def test_stream_endpoint_sends_failed_safe_error(monkeypatch):
    store = InMemoryJobStore()
    store.create_job(
        _job(
            "job-stream-failed",
            status=JobStatus.FAILED,
            progress=100,
            error="failed to process DOCX file",
        )
    )
    monkeypatch.setattr("app.api.stream.get_job_store", lambda: store)
    monkeypatch.setattr(
        "app.api.stream.get_progress_event_store",
        lambda: _StreamEventStore([]),
    )

    response = TestClient(app).get("/stream/job-stream-failed")

    assert response.status_code == 200
    assert '"stage":"failed"' in response.text
    assert "failed to process DOCX file" in response.text
    assert "Traceback" not in response.text


def test_stream_endpoint_events_do_not_include_document_text(monkeypatch):
    store = InMemoryJobStore()
    store.create_job(
        _job(
            "job-stream-history",
            status=JobStatus.COMPLETED,
            progress=100,
        )
    )
    events = [
        build_progress_event(
            job_id="job-stream-history",
            stage=JobStatus.TRANSLATING,
            progress=55,
            message="Batch 1/2 translated",
        ),
        build_progress_event(
            job_id="job-stream-history",
            stage=JobStatus.COMPLETED,
            progress=100,
            message="Completed",
        ),
    ]
    monkeypatch.setattr("app.api.stream.get_job_store", lambda: store)
    monkeypatch.setattr(
        "app.api.stream.get_progress_event_store",
        lambda: _StreamEventStore(events),
    )

    response = TestClient(app).get("/stream/job-stream-history")

    assert response.status_code == 200
    assert "Sensitive source paragraph" not in response.text
    assert "Batch 1/2 translated" in response.text


def test_worker_processes_translation_job_with_mock_ai(tmp_path, monkeypatch):
    store = InMemoryJobStore()
    source_path = _create_docx(tmp_path, "worker.docx", "Hello")
    job = TranslationJob(
        job_id="job-worker",
        source_lang=LanguageCode.EN,
        target_lang=LanguageCode.RU,
        original_filename="worker.docx",
        upload_path=source_path.as_posix(),
    )
    store.create_job(job)

    monkeypatch.setattr("workers.translation_worker.get_job_store", lambda: store)
    monkeypatch.setattr(
        "workers.translation_worker.get_settings",
        lambda: Settings(
            mock_ai_enabled=True,
            upload_dir=tmp_path / "uploads",
            output_dir=tmp_path / "outputs",
            tmp_dir=tmp_path / "tmp",
            translation_memory_db_path=tmp_path / "memory.sqlite3",
        ),
    )

    process_translation_job("job-worker")

    completed_job = store.get_job("job-worker")
    assert completed_job is not None
    assert completed_job.status == JobStatus.COMPLETED
    assert completed_job.progress == 100
    assert completed_job.result_file is not None

    translated_document = Document(completed_job.result_file)
    assert translated_document.paragraphs[0].text == "Hello [ru]"


def test_worker_failed_job_stores_safe_error(tmp_path, monkeypatch):
    store = InMemoryJobStore()
    job = TranslationJob(
        job_id="job-failed",
        source_lang=LanguageCode.EN,
        target_lang=LanguageCode.RU,
        original_filename="missing.docx",
        upload_path=(tmp_path / "missing.docx").as_posix(),
    )
    store.create_job(job)

    monkeypatch.setattr("workers.translation_worker.get_job_store", lambda: store)
    monkeypatch.setattr(
        "workers.translation_worker.get_settings",
        lambda: Settings(
            mock_ai_enabled=True,
            upload_dir=tmp_path / "uploads",
            output_dir=tmp_path / "outputs",
            tmp_dir=tmp_path / "tmp",
            translation_memory_db_path=tmp_path / "memory.sqlite3",
        ),
    )

    with pytest.raises(DocumentProcessingError):
        process_translation_job("job-failed")

    failed_job = store.get_job("job-failed")
    assert failed_job is not None
    assert failed_job.status == JobStatus.FAILED
    assert failed_job.progress == 100
    assert failed_job.error == "failed to process DOCX file"
    assert "Traceback" not in failed_job.error


def _block(block_id: str, text: str, translatable: bool = True) -> DocumentBlock:
    return DocumentBlock(
        block_id=block_id,
        text=text,
        location=f"test:{block_id}",
        translatable=translatable,
    )


def _job(
    job_id: str,
    status: JobStatus = JobStatus.QUEUED,
    progress: int = 0,
    error: str | None = None,
    result_file: str | None = None,
) -> TranslationJob:
    return TranslationJob(
        job_id=job_id,
        status=status,
        progress=progress,
        source_lang=LanguageCode.EN,
        target_lang=LanguageCode.RU,
        original_filename="source.docx",
        upload_path="uploads/source.docx",
        result_file=result_file,
        error=error,
    )


def _create_docx(tmp_path, filename: str, text: str):
    path = tmp_path / filename
    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return path


class _UploadFileStub:
    def __init__(self, filename: str, content_type: str, content: bytes) -> None:
        self.filename = filename
        self.content_type = content_type
        self._content = content

    async def read(self) -> bytes:
        return self._content


class _FakeRedisClient:
    def __init__(self) -> None:
        self.values: dict[str, str] = {}
        self.set_calls: list[tuple[str | None, int, str]] = []

    def get(self, name: str) -> str | None:
        return self.values.get(name)

    def set(self, name: str | None, value: str, ex: int | None = None) -> None:
        self.set_calls.append((name, ex or 0, value))
        if name is not None:
            self.values[name] = value


class _FailingRedisClient:
    def get(self, name: str) -> str | None:
        raise RedisError("redis unavailable")

    def set(self, name: str, value: str, ex: int | None = None) -> None:
        raise RedisError("redis unavailable")


class _PipelineCache:
    def __init__(self, translations_by_text: dict[str, str]) -> None:
        self.translations_by_key = {
            build_cache_key(text, LanguageCode.EN, LanguageCode.RU): translation
            for text, translation in translations_by_text.items()
        }
        self.stored: list[tuple[str | None, str]] = []

    def get_translation(self, key: str | None) -> str | None:
        return self.translations_by_key.get(key)

    def set_translation(
        self,
        key: str | None,
        value: str,
        ttl: int | None = None,
    ) -> None:
        self.stored.append((key, value))


class _FakeTranslationClient:
    def __init__(self, translations: dict[str, str]) -> None:
        self.translations = translations
        self.requested_block_ids: list[list[str]] = []
        self.requested_glossary_terms: list[list[tuple[str, str]]] = []

    async def translate_blocks(
        self,
        blocks: list[DocumentBlock],
        source_lang: LanguageCode,
        target_lang: LanguageCode,
        glossary_terms=None,
    ) -> dict[str, str]:
        self.requested_block_ids.append([block.block_id for block in blocks])
        self.requested_glossary_terms.append(
            [
                (term.source, term.target)
                for term in (glossary_terms or [])
            ]
        )
        return self.translations


class _StreamEventStore:
    def __init__(self, events) -> None:
        self.events = events

    def history_with_ids(self, job_id: str):
        return [
            (f"{index}-0", event)
            for index, event in enumerate(self.events, start=1)
            if event.job_id == job_id
        ]

    def read_after(self, job_id: str, last_event_id: str, block_ms: int):
        return last_event_id, []


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)
